"""Generate Word document explaining poster Sections 2 & 3 in detail."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Poster Sections 2 & 3 — Detailed Explanation', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Reference guide for poster presentation and Q&A')
doc.add_paragraph('')

# ============= SECTION 2 =============
doc.add_heading('Section 2: DQN & Double DQN', level=1)

doc.add_heading('CNN Architecture', level=2)
doc.add_paragraph(
    'The neural network takes the board state (12x5x5 binary tensor) as input '
    'and outputs 625 Q-values, one for each possible action (from-square x to-square).'
)

doc.add_heading('Layer-by-layer breakdown:', level=3)
t = doc.add_table(rows=7, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Layer'
t.rows[0].cells[1].text = 'What it does'
data = [
    ('Conv2d(12->64, 3x3)', 'Scans the board with 64 different 3x3 filters. Detects local patterns like "two pieces next to each other" or "pawn in front of empty square".'),
    ('Conv2d(64->128, 3x3)', '128 filters combine features from layer 1. Detects higher-level patterns like "bishop has open diagonal" or "rook controls a file".'),
    ('Conv2d(128->128, 3x3)', 'Further refinement. Captures complex spatial relationships like "queen can attack the king through a diagonal".'),
    ('BN (BatchNorm)', 'Normalizes output of each layer so training is more stable. Prevents values from exploding or vanishing.'),
    ('ReLU', 'Activation function: max(0, x). Allows the network to learn non-linear relationships. Without it, stacking layers would be useless.'),
    ('FC(3200->512)->FC(512->625)', 'Flatten 128x5x5=3200 features into a vector, compress to 512, then output 625 Q-values (one per action).'),
]
for i, (layer, desc) in enumerate(data):
    t.rows[i+1].cells[0].text = layer
    t.rows[i+1].cells[1].text = desc

doc.add_paragraph('')
doc.add_heading('DQN Target Computation', level=2)
p = doc.add_paragraph()
p.add_run('y = r + gamma * max_a\' Q(s\', a\'; theta-)').bold = True
doc.add_paragraph(
    'In plain English: the training target y equals the immediate reward (r) '
    'plus the discounted (gamma=0.99) maximum Q-value of the next state, '
    'estimated by the old/target network (theta-).\n\n'
    'The network\'s current estimate Q(s,a;theta) is compared against this target. '
    'If they differ, the weights are adjusted via gradient descent to reduce the gap. '
    'This is the same idea as tabular Q-learning: '
    'Q(s,a) <- Q(s,a) + alpha[r + gamma*maxQ(s\',a\') - Q(s,a)], '
    'but using a neural network instead of a table.'
)

doc.add_paragraph('')
doc.add_heading('Double DQN Target Computation', level=2)
p = doc.add_paragraph()
p.add_run('a* = argmax_a\' Q(s\', a\'; theta)     <- online network selects action\n'
          'y  = r + gamma * Q(s\', a*; theta-)    <- target network evaluates it').bold = True

doc.add_paragraph(
    '\nThe key difference from standard DQN: action selection and value evaluation are decoupled.'
)

t2 = doc.add_table(rows=3, cols=3)
t2.style = 'Table Grid'
t2.rows[0].cells[0].text = ''
t2.rows[0].cells[1].text = 'Who selects action?'
t2.rows[0].cells[2].text = 'Who evaluates it?'
t2.rows[1].cells[0].text = 'DQN'
t2.rows[1].cells[1].text = 'Target network (theta-)'
t2.rows[1].cells[2].text = 'Target network (theta-) -- SAME'
t2.rows[2].cells[0].text = 'Double DQN'
t2.rows[2].cells[1].text = 'Online network (theta)'
t2.rows[2].cells[2].text = 'Target network (theta-) -- DIFFERENT'

doc.add_paragraph(
    '\nWhy does this matter? When DQN uses the same network for both, '
    'it always picks the action with the highest (possibly noisy) Q-value. '
    'If one action\'s Q-value is overestimated due to noise, the max operator will select it, '
    'causing systematic upward bias. Double DQN fixes this by using one network to pick '
    'and a different network to score -- reducing overestimation.\n\n'
    'Our data confirms this: DQN avg Q = 1.891, Double DQN avg Q = 1.841. '
    'DQN overestimates by about 13% more than Double DQN.'
)

doc.add_paragraph('')
doc.add_heading('Key Hyperparameters Explained', level=2)

t3 = doc.add_table(rows=8, cols=3)
t3.style = 'Table Grid'
t3.rows[0].cells[0].text = 'Parameter'
t3.rows[0].cells[1].text = 'Value'
t3.rows[0].cells[2].text = 'What it means'
hdata = [
    ('Learning rate', '1e-4', 'How big each weight update step is. Too large = unstable. Too small = learns too slowly.'),
    ('Replay buffer', '100K transitions', 'Stores the most recent 100,000 single-step experiences (s,a,r,s\'). NOT 100K games -- about 2000-2500 games worth. Each training step randomly samples 256 from this buffer.'),
    ('Target update', 'Every 1000 steps', 'Target network theta- is synced with online network theta every 1000 gradient steps. Provides stable regression target.'),
    ('Epsilon-greedy', '1.0 -> 0.05', 'Starts at 100% random moves, linearly decays to 5% random over 30K episodes. At 20K episodes it was still ~0.37 (37% random).'),
    ('Discount gamma', '0.99', 'How much future rewards are valued. A reward 50 steps away is worth 0.99^50 = 0.61 of its face value.'),
    ('Huber Loss', 'Smooth L1', 'More robust than MSE. Acts like MSE for small errors but L1 for large errors, preventing gradient explosion.'),
    ('Q-value clamp', '[-2, 2]', 'Target Q-values clamped to this range to prevent divergence. Since rewards are only +/-1, Q should not exceed +/-2.'),
]
for i, (p, v, d) in enumerate(hdata):
    t3.rows[i+1].cells[0].text = p
    t3.rows[i+1].cells[1].text = v
    t3.rows[i+1].cells[2].text = d

doc.add_paragraph('')

# ============= SECTION 3 =============
doc.add_heading('Section 3: MCTS (Planning)', level=1)

doc.add_heading('Core Idea', level=2)
doc.add_paragraph(
    'MCTS is fundamentally different from DQN. It does NOT learn or train. '
    'Every time it needs to make a move, it runs 200 simulated games from the current position '
    'to estimate which move is best. This is a planning/search method, not a learning method.'
)

doc.add_paragraph('')
doc.add_heading('The Four Phases (per simulation)', level=2)

t4 = doc.add_table(rows=5, cols=3)
t4.style = 'Table Grid'
t4.rows[0].cells[0].text = 'Phase'
t4.rows[0].cells[1].text = 'Code'
t4.rows[0].cells[2].text = 'Explanation'
mdata = [
    ('1. Selection', 'SELECT(root) // UCB1', 'From root, walk down tree picking child with highest UCB1 score. Balances promising moves (exploitation) and under-explored moves (exploration).'),
    ('2. Expansion', 'EXPAND(node)', 'At a node with untried moves, pick one and create a new child. Grows search tree one node at a time.'),
    ('3. Rollout', 'ROLLOUT(node) // random', 'From the new node, both sides play completely random legal moves until game ends. Record result: +1 white wins, -1 black wins, 0 draw.'),
    ('4. Backpropagation', 'BACKPROP(node, v)', 'Pass result back up through every node in this path. Each node updates visit count N and total value W.'),
]
for i, (phase, code, desc) in enumerate(mdata):
    t4.rows[i+1].cells[0].text = phase
    t4.rows[i+1].cells[1].text = code
    t4.rows[i+1].cells[2].text = desc

doc.add_paragraph('')
doc.add_heading('Concrete Example', level=2)
doc.add_paragraph('Suppose current position has 3 legal moves: A, B, C.')
examples = [
    'Sim 1: Try A -> random play to end -> White wins (+1)  -> A: 1 win / 1 visit',
    'Sim 2: Try B -> random play to end -> Black wins (-1)  -> B: 0 wins / 1 visit',
    'Sim 3: Try C -> random play to end -> White wins (+1)  -> C: 1 win / 1 visit',
    'Sim 4: Try A -> random play to end -> White wins (+1)  -> A: 2 wins / 2 visits',
    '... (repeat 200 times) ...',
    'Final: A visited 80 times, B visited 50, C visited 70',
    '-> Choose A (most visited = most confident)',
]
for ex in examples:
    doc.add_paragraph(ex, style='List Bullet')

doc.add_paragraph('')
doc.add_heading('UCB1 Formula', level=2)
p = doc.add_paragraph()
p.add_run('score = Q/N + c * sqrt(ln(N_parent) / N)').bold = True

t5 = doc.add_table(rows=4, cols=2)
t5.style = 'Table Grid'
t5.rows[0].cells[0].text = 'Component'
t5.rows[0].cells[1].text = 'Meaning'
t5.rows[1].cells[0].text = 'Q/N (exploitation)'
t5.rows[1].cells[1].text = 'Average win rate of this move. High = usually wins.'
t5.rows[2].cells[0].text = 'c*sqrt(ln(N_parent)/N) (exploration)'
t5.rows[2].cells[1].text = 'Bonus for under-explored moves. Few visits -> large bonus -> try it more.'
t5.rows[3].cells[0].text = 'c = sqrt(2) = 1.414'
t5.rows[3].cells[1].text = 'Exploration constant. Larger = more exploration, smaller = more exploitation.'

doc.add_paragraph('')
doc.add_heading('Value from White\'s Perspective', level=2)
doc.add_paragraph(
    'All nodes store values from White\'s viewpoint: White win = +1, Black win = -1, Draw = 0. '
    'During selection, White nodes maximize UCB1 (best for White), '
    'Black nodes minimize it (best for Black = worst for White). '
    'This avoids flipping signs at every tree level.'
)

doc.add_paragraph('')

# ============= COMPARISON =============
doc.add_heading('DQN vs MCTS: Key Comparison', level=1)
t6 = doc.add_table(rows=6, cols=3)
t6.style = 'Table Grid'
t6.rows[0].cells[0].text = 'Aspect'
t6.rows[0].cells[1].text = 'DQN / Double DQN'
t6.rows[0].cells[2].text = 'MCTS'
cdata = [
    ('Requires training?', 'Yes - 20K games, ~24 hours', 'No - works immediately'),
    ('Decision speed', 'Fast - one CNN forward pass', 'Slow - 200 simulations per move'),
    ('How it picks moves', 'CNN estimates Q-value per action, picks highest', 'Simulates many random games, picks most-visited'),
    ('Handles sparse rewards?', 'Poorly - must bootstrap signal back 40-50 steps', 'Well - simulates directly to game end'),
    ('Win rate vs Random', 'DQN 13%, DDQN 18%', 'MCTS 93%'),
]
for i, (a, d, m) in enumerate(cdata):
    t6.rows[i+1].cells[0].text = a
    t6.rows[i+1].cells[1].text = d
    t6.rows[i+1].cells[2].text = m

out = 'poster/poster_sections_explained.docx'
doc.save(out)
print(f'Saved: {out}')
