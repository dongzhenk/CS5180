"""Generate Word documents for presentation script and Q&A prep."""
from docx import Document
from docx.shared import Pt, RGBColor

# ========================================
# 1. Presentation Script
# ========================================
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

doc.add_heading('Poster Presentation Script (~5 min)', level=0)

sections = [
    ('Opening (30s)',
     'Hi, I\'m Zhenkang. My project compares three approaches for playing '
     '5x5 Gardner Mini-Chess: DQN, Double DQN, and MCTS. The goal is to study '
     'how value-based reinforcement learning behaves in a two-player adversarial '
     'setting -- specifically convergence stability and overestimation bias.'),

    ('1. Problem -- Why Mini-Chess? (40s)',
     '[Point to Section 1]\n\n'
     'Gardner Mini-Chess is a 5x5 variant with 10 pieces per side. It\'s '
     'computationally manageable but strategically rich -- it\'s been weakly '
     'solved and the game-theoretic value is a draw.\n\n'
     'I formulated it as an MDP: the state is a 12-channel binary encoding -- '
     '6 channels for my pieces, 6 for the opponent. The action space is 625 '
     'from-square-to-square combinations with legal-action masking. Rewards are '
     'terminal only -- +1 for a win, -1 for a loss, 0 for a draw.'),

    ('2. Algorithms (60s)',
     '[Point to Sections 2 & 3]\n\n'
     'For DQN, I use a 3-layer CNN that outputs Q-values for all 625 actions. '
     'I train with experience replay and a target network.\n\n'
     'Double DQN uses the exact same architecture -- the only difference is '
     'the target computation. Standard DQN uses the target network for both '
     'action selection and evaluation, which causes overestimation. Double DQN '
     'decouples these: the online network selects, the target network evaluates.\n\n'
     'MCTS is our planning baseline. UCB1 selection, random rollouts, 200 '
     'simulations per move. No training needed.'),

    ('3. Training Setup (30s)',
     '[Point to Section 4]\n\n'
     'Training uses frozen-opponent self-play. The agent plays against a frozen '
     'copy of itself, updated every 500 episodes. Canonical state encoding lets '
     'one network play both colors. 20,000 episodes, 2 random seeds.'),

    ('4. Results -- Key Findings (90s)',
     '[Point to Sections 5-7]\n\n'
     'Training win rate: both peak ~25% during warmup, then stabilize ~15% '
     'as the frozen opponent gets stronger.\n\n'
     '[Point to Q-value figure]\n\n'
     'The Q-value analysis is the core. DQN avg Q = 1.89, DDQN = 1.84. '
     'DQN overestimates more. DDQN has lower cross-seed variance '
     '(+/-0.8% vs +/-1.4%).\n\n'
     '[Point to evaluation bar chart]\n\n'
     'MCTS: 93% win vs Random, dominates. DQN: 13%, DDQN: 18%. Against '
     'heuristic, MCTS 63%, DQN/DDQN near 0-5%. High draw rate means agents '
     'learned not to lose but not to win -- sparse rewards are the bottleneck.'),

    ('5. Conclusions (30s)',
     '[Point to Section 8]\n\n'
     'Three takeaways: (1) MCTS >> DQN/DDQN -- planning beats learning under '
     'sparse rewards. (2) Double DQN shows lower overestimation, confirming '
     'the theory. (3) Sparse rewards cause conservative play. Future: reward '
     'shaping, AlphaZero-style MCTS+NN.\n\n'
     'Thanks -- happy to take questions.'),
]

for title, content in sections:
    doc.add_heading(title, level=1)
    for para in content.split('\n\n'):
        p = doc.add_paragraph(para.strip())
        if para.strip().startswith('['):
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)

doc.save('poster/presentation_script.docx')
print('Saved: poster/presentation_script.docx')


# ========================================
# 2. Q&A Prep Bilingual
# ========================================
doc2 = Document()
s2 = doc2.styles['Normal']
s2.font.name = 'Calibri'
s2.font.size = Pt(10.5)

doc2.add_heading('Q&A Preparation (Bilingual)', level=0)

# ---- Part 1 ----
doc2.add_heading('Part 1: Likely Questions / \u53ef\u80fd\u88ab\u95ee\u7684\u95ee\u9898', level=1)

qa = [
    ('Q1: Why 5x5 instead of 8x8?',
     '\u4e3a\u4ec0\u4e48\u90095x5\u800c\u4e0d\u662f8x8\uff1f',
     'Computational feasibility. Standard chess ~10^47 states. 5x5 has 625 actions, feasible on single GPU. Still strategically rich (weakly solved, value = draw).',
     '\u8ba1\u7b97\u53ef\u884c\u6027\u3002\u6807\u51c6\u56fd\u9645\u8c61\u68cb~10^47\u72b6\u6001\u30025x5\u53ea\u6709625\u4e2a\u52a8\u4f5c\uff0c\u5355GPU\u53ef\u8bad\u7ec3\u3002\u4ecd\u7136\u7b56\u7565\u4e30\u5bcc\uff08\u5f31\u89e3\u503c=\u5e73\u5c40\uff09\u3002'),

    ('Q2: Why is win rate so low (13-18%)?',
     '\u80dc\u7387\u4e3a\u4ec0\u4e48\u8fd9\u4e48\u4f4e\uff1f',
     '(1) Sparse terminal rewards -- feedback only at game end, 40-50 moves away. (2) Self-play instability. High draw rate = agents learned not to lose, not to win.',
     '(1) \u7a00\u758f\u7ec8\u5c40\u5956\u52b1\u2014\u201440-50\u6b65\u540e\u624d\u6709\u53cd\u9988\u3002(2) \u81ea\u5bf9\u5f08\u4e0d\u7a33\u5b9a\u3002\u9ad8\u5e73\u5c40\u7387=\u5b66\u4f1a\u4e86\u4e0d\u8f93\uff0c\u6ca1\u5b66\u4f1a\u8d62\u3002'),

    ('Q3: Q-value overestimation significance?',
     'Q\u503c\u8fc7\u9ad8\u4f30\u8ba1\u8bf4\u660e\u4ec0\u4e48\uff1f',
     'Core finding. DQN max operator picks noisiest high estimate = positive bias. DDQN decouples selection/evaluation. DQN Q=1.89 > DDQN 1.84. Confirms Van Hasselt 2016.',
     '\u6838\u5fc3\u53d1\u73b0\u3002DQN\u7684max\u64cd\u4f5c\u9009\u6700\u4e50\u89c2\u4f30\u8ba1=\u6b63\u504f\u5dee\u3002DDQN\u89e3\u8026\u9009\u62e9/\u8bc4\u4f30\u3002DQN Q=1.89 > DDQN 1.84\u3002\u9a8c\u8bc1Van Hasselt 2016\u3002'),

    ('Q4: Why does MCTS dominate?',
     '\u4e3a\u4ec0\u4e48MCTS\u78be\u538b\uff1f',
     'Direct lookahead search, simulates future. No training needed. DQN must learn from experience. In small deterministic games, search > learning. AlphaZero combined both.',
     '\u76f4\u63a5\u524d\u77bb\u641c\u7d22\uff0c\u6a21\u62df\u672a\u6765\u3002\u4e0d\u9700\u8bad\u7ec3\u3002DQN\u9700\u4ece\u7ecf\u9a8c\u5b66\u4e60\u3002\u5c0f\u578b\u786e\u5b9a\u6027\u535a\u5f08\u4e2d\u641c\u7d22>\u5b66\u4e60\u3002AlphaZero\u7ed3\u5408\u4e86\u4e24\u8005\u3002'),

    ('Q5: Frozen-opponent self-play?',
     '\u51bb\u7ed3\u5bf9\u624b\u81ea\u5bf9\u5f08\uff1f',
     'Latest-version opponent = unstable (violates stationary MDP). Frozen copy = stable signal, refreshed every 500 eps. Like fictitious self-play.',
     '\u5bf9\u6700\u65b0\u7248\u672c=\u4e0d\u7a33\u5b9a\uff08\u8fdd\u53cdMDP\u5047\u8bbe\uff09\u3002\u51bb\u7ed3\u526f\u672c=\u7a33\u5b9a\u4fe1\u53f7\uff0c\u6bcf500\u5c40\u66f4\u65b0\u3002\u7c7b\u4f3c\u865a\u62df\u81ea\u5bf9\u5f08\u3002'),

    ('Q6: Would more training help?',
     '\u66f4\u591a\u8bad\u7ec3\u6709\u7528\u5417\uff1f',
     'Some improvement, diminishing returns. Sparse reward bottleneck remains. Reward shaping would help more.',
     '\u4f1a\u6709\u4e00\u4e9b\u63d0\u5347\uff0c\u4f46\u6536\u76ca\u9012\u51cf\u3002\u7a00\u758f\u5956\u52b1\u74f6\u9888\u4ecd\u5728\u3002\u5956\u52b1\u5851\u5f62\u66f4\u6709\u6548\u3002'),

    ('Q7: Why Huber loss not MSE?',
     '\u4e3a\u4ec0\u4e48\u7528Huber\u4e0d\u7528MSE\uff1f',
     'MSE caused loss explosion (10^6). Huber = robust to outliers. Also Q-value clamped to [-2,2].',
     'MSE\u5bfc\u81f4loss\u7206\u70b8(10^6)\u3002Huber\u5bf9\u5f02\u5e38\u503c\u9c81\u68d2\u3002\u540c\u65f6Q\u503c\u88c1\u526a\u5230[-2,2]\u3002'),

    ('Q8: Canonical encoding?',
     '\u89c4\u8303\u7f16\u7801\u600e\u4e48\u5de5\u4f5c\uff1f',
     'State always from current player\'s view. Black\'s turn = rotate 180 + swap colors. One network plays both sides.',
     '\u72b6\u6001\u6c38\u8fdc\u4ece\u5f53\u524d\u73a9\u5bb6\u89c6\u89d2\u3002\u9ed1\u65b9\u65f6\u65cb\u8f6c180\u00b0+\u4ea4\u6362\u989c\u8272\u3002\u4e00\u4e2a\u7f51\u7edc\u4e24\u4e2a\u989c\u8272\u90fd\u80fd\u4e0b\u3002'),

    ('Q9: Only 2 seeds enough?',
     '\u53ea\u67092\u4e2aseed\u591f\u5417\uff1f',
     'Limited, I acknowledge. Results consistent across seeds. Each ~6h on RTX 4060. Plan more seeds if time permits.',
     '\u6709\u9650\uff0c\u6211\u627f\u8ba4\u3002\u7ed3\u679c\u5728seed\u95f4\u4e00\u81f4\u3002\u6bcf\u4e2a~6h\u3002\u65f6\u95f4\u5141\u8bb8\u4f1a\u8ddf\u591a\u3002'),

    ('Q10: What is the heuristic opponent?',
     '\u542f\u53d1\u5f0f\u5bf9\u624b\u662f\u4ec0\u4e48\uff1f',
     'Hand-coded rules: check > capture > random. Stronger than random (exploits tactics) but no strategy.',
     '\u624b\u5199\u89c4\u5219\uff1a\u5c06\u519b>\u5403\u5b50>\u968f\u673a\u3002\u6bd4\u968f\u673a\u5f3a\uff08\u4f1a\u5229\u7528\u6218\u672f\uff09\u4f46\u6ca1\u6709\u6218\u7565\u3002'),
]

for q_en, q_cn, a_en, a_cn in qa:
    doc2.add_heading(q_en, level=2)
    p = doc2.add_paragraph(q_cn)
    p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    doc2.add_paragraph('EN: ' + a_en)
    p2 = doc2.add_paragraph('CN: ' + a_cn)
    p2.runs[0].font.color.rgb = RGBColor(0, 0, 160)

# ---- Part 2 ----
doc2.add_heading('Part 2: Ask the Professor / \u95ee\u6559\u6388\u7684\u95ee\u9898', level=1)

prof = [
    ('Q1: Would reward shaping help or introduce bias?',
     '\u5956\u52b1\u5851\u5f62\u4f1a\u6709\u5e2e\u52a9\u8fd8\u662f\u5f15\u5165\u504f\u5dee\uff1f',
     'Why: shows you think about solutions.',
     'Expected: potential-based shaping preserves optimal policy. Capture rewards common but may bias short-term.'),

    ('Q2: Would PPO be more suitable?',
     'PPO\u662f\u5426\u66f4\u9002\u5408\uff1f',
     'Why: shows breadth.',
     'Expected: PPO more stable but on-policy. AlphaZero = MCTS + policy-value net.'),

    ('Q3: Why so many draws theoretically?',
     '\u5927\u91cf\u5e73\u5c40\u7684\u7406\u8bba\u89e3\u91ca\uff1f',
     'Why: shows analytical depth.',
     'Expected: gamma=0.99 discounts too much over 50 steps. Try gamma=1.0.'),

    ('Q4: Simplest way to add MCTS+DQN like AlphaZero?',
     '\u6700\u7b80\u5355\u7684MCTS+DQN\u7ed3\u5408\u65b9\u5f0f\uff1f',
     'Why: shows ambition.',
     'Expected: use Q-net as rollout policy, or add policy head + MCTS targets.'),
]

for q_en, q_cn, why, expected in prof:
    doc2.add_heading(q_en, level=2)
    p = doc2.add_paragraph(q_cn)
    p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    pi = doc2.add_paragraph(why)
    pi.runs[0].italic = True
    doc2.add_paragraph(expected)

# ---- Part 3 ----
doc2.add_heading('Part 3: Tough Questions / \u68d8\u624b\u95ee\u9898', level=1)

tough = [
    ('"DQN barely beats random. Failure?"',
     'Limited result, not failure. Goal = compare algorithms, not solve chess. Confirmed DDQN theory. Low performance itself is an interesting finding about sparse rewards.',
     '\u53d7\u9650\u7ed3\u679c\uff0c\u4e0d\u662f\u5931\u8d25\u3002\u76ee\u6807=\u6bd4\u8f83\u7b97\u6cd5\u884c\u4e3a\u3002\u9a8c\u8bc1\u4e86DDQN\u7406\u8bba\u3002\u4f4e\u6027\u80fd\u672c\u8eab\u662f\u5173\u4e8e\u7a00\u758f\u5956\u52b1\u7684\u6709\u8da3\u53d1\u73b0\u3002'),

    ('"2 seeds not significant."',
     'Correct, limited power. Consistent results encouraging. ~6h per seed, practical tradeoff.',
     '\u6b63\u786e\uff0c\u7edf\u8ba1\u80fd\u529b\u6709\u9650\u3002\u7ed3\u679c\u4e00\u81f4\u6027\u4ee4\u4eba\u9f13\u821e\u3002\u6bcf seed~6h\uff0c\u5b9e\u9645\u6743\u8861\u3002'),

    ('"Why not AlphaZero?"',
     'Different paradigm (MCTS+NN). This project isolates components: pure learning (DQN), improved learning (DDQN), pure planning (MCTS). Understanding parts before combining.',
     '\u4e0d\u540c\u8303\u5f0f\u3002\u672c\u9879\u76ee\u9694\u79bb\u5404\u7ec4\u4ef6\uff1a\u7eaf\u5b66\u4e60(DQN)\u3001\u6539\u8fdb\u5b66\u4e60(DDQN)\u3001\u7eaf\u89c4\u5212(MCTS)\u3002\u7ec4\u5408\u524d\u5148\u7406\u89e3\u5404\u90e8\u5206\u3002'),
]

for q, a_en, a_cn in tough:
    doc2.add_heading(q, level=2)
    doc2.add_paragraph('EN: ' + a_en)
    p = doc2.add_paragraph('CN: ' + a_cn)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 160)

doc2.save('poster/qa_prep_bilingual.docx')
print('Saved: poster/qa_prep_bilingual.docx')
